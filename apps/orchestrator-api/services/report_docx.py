"""
report_docx — render a report's Markdown body into a Microsoft Word (.docx)
file in memory (bytes), so it can be emailed as an attachment.

Handles the subset of Markdown the reports use:
  - #/##/### headings        → Word heading styles
  - | a | b | tables          → a real Word table (header row bold)
  - - bullets                 → List Bullet paragraphs
  - **bold** inline           → bold runs

Pure python-docx (already in requirements). Returns bytes; never raises to the
caller — on any failure it falls back to a plain-text docx of the raw body.
"""

from __future__ import annotations

import io
import re

from services.logger import log


_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a real clickable hyperlink run to a paragraph (blue, underlined)."""
    try:
        from docx.oxml.shared import OxmlElement, qn
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
        und = OxmlElement("w:u"); und.set(qn("w:val"), "single"); rpr.append(und)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        link.append(run)
        paragraph._p.append(link)
    except Exception:
        paragraph.add_run(f"{text} ({url})")


def _add_bold_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** spans and [text](url) links
    (rendered as real clickable hyperlinks)."""
    pos = 0
    for m in _LINK_RE.finditer(text):
        _add_bold_runs(paragraph, text[pos:m.start()])
        _add_hyperlink(paragraph, m.group(2), m.group(1))
        pos = m.end()
    _add_bold_runs(paragraph, text[pos:])


def _add_table(doc, block: list[str]) -> None:
    rows: list[list[str]] = []
    for ln in block:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if cells and all(set(c) <= set("-: ") for c in cells):
            continue  # separator row
        rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncol)
    for style in ("Light Grid Accent 1", "Light List Accent 1", "Table Grid"):
        try:
            table.style = style
            break
        except Exception:
            continue
    for ri, r in enumerate(rows):
        for ci, val in enumerate(r):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            if ri == 0:
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.bold = True


def markdown_to_docx(md: str, title: str, subtitle: str = "") -> bytes:
    """Render Markdown → .docx bytes."""
    from docx import Document

    try:
        doc = Document()
        doc.add_heading(title, level=0)
        if subtitle:
            p = doc.add_paragraph(subtitle)
            if p.runs:
                p.runs[0].italic = True

        lines = (md or "").split("\n")
        i, n = 0, len(lines)
        while i < n:
            line = lines[i]
            is_tbl = (line.lstrip().startswith("|") and i + 1 < n
                      and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"))
            if is_tbl:
                block = []
                while i < n and lines[i].lstrip().startswith("|"):
                    block.append(lines[i]); i += 1
                _add_table(doc, block)
                continue
            m = re.match(r"^(#{1,6})\s*(.*)$", line)
            if m:
                doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
                i += 1
                continue
            if re.match(r"^\s*[-*]\s+", line):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
                i += 1
                continue
            if line.strip() == "":
                i += 1
                continue
            _add_runs(doc.add_paragraph(), line)
            i += 1

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
    except Exception as e:
        log.warning(f"report_docx: render failed, using plain fallback: {e}")
        doc = Document()
        doc.add_heading(title, level=0)
        for para in (md or "").split("\n\n"):
            doc.add_paragraph(para)
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
