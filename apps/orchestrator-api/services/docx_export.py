"""
VIP AI Platform — Markdown → DOCX Exporter
Converts twin task results (Markdown) into Microsoft Word .docx for download.
Uses python-docx (already installed).
"""

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(
    title: str,
    markdown_text: str,
    author: str = "Davronbek's Twin",
    subtitle: str = "",
) -> bytes:
    """
    Render a (lightweight) Markdown document into a styled .docx and return bytes.
    Supports: # / ## / ### headings, bold (**text**), italic (*text*), bullets (- ),
    numbered lists, code blocks (```), simple tables (| col | col |), horizontal rules (---).
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # === Cover header ===
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)

    if subtitle:
        sub = doc.add_paragraph()
        srun = sub.add_run(subtitle)
        srun.italic = True
        srun.font.size = Pt(11)
        srun.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    meta = doc.add_paragraph()
    mrun = meta.add_run(f"Author: {author}    ·    Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    mrun.font.size = Pt(9)
    mrun.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    doc.add_paragraph().add_run("").font.size = Pt(2)  # spacer

    # === Body parser ===
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    i = 0
    in_code_block = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_code():
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        code_buf.clear()

    def flush_table():
        if not table_buf:
            return
        # Skip the separator row that markdown tables use (---|---)
        cleaned = [r for r in table_buf if not all(set(c.strip()) <= set("-:| ") for c in r)]
        if not cleaned:
            table_buf.clear()
            return
        cols = max(len(r) for r in cleaned)
        t = doc.add_table(rows=len(cleaned), cols=cols)
        t.style = "Light Grid Accent 1"
        for r_i, row in enumerate(cleaned):
            for c_i in range(cols):
                cell_text = row[c_i] if c_i < len(row) else ""
                cell = t.cell(r_i, c_i)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(10)
                if r_i == 0:
                    run.bold = True
        doc.add_paragraph().add_run("").font.size = Pt(4)
        table_buf.clear()

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buf.append(line)
            i += 1
            continue

        # Tables (rows that contain | )
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        elif table_buf:
            flush_table()

        # Headings
        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True; r.font.size = Pt(13)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True; r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.bold = True; r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 50)
        # Bullet list
        elif line.lstrip().startswith(("- ", "* ", "+ ")):
            stripped = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped)
        # Numbered list
        elif re.match(r"^\s*\d+\.\s", line):
            stripped = re.sub(r"^\s*\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, stripped)
        # Empty line
        elif not line.strip():
            doc.add_paragraph()
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)
        i += 1

    flush_code()
    flush_table()

    # === Footer ===
    doc.add_paragraph()
    foot = doc.add_paragraph()
    frun = foot.add_run("Generated by VIP AI Platform — Davronbek's Twin")
    frun.font.size = Pt(8)
    frun.italic = True
    frun.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_inline(paragraph, text: str):
    """Parse simple inline Markdown: **bold**, *italic*, `code`."""
    # Tokenize
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()]).font.size = Pt(11)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True; r.font.size = Pt(11)
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True; r.font.size = Pt(11)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:]).font.size = Pt(11)
