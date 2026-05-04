"""
VIP AI Platform — File Text Extraction
Extracts plain text from uploaded files so the Twin can read them in chat.

Supported:
- Plain text: .txt, .md, .csv, .json, .log, .py, .js, .ts, etc.
- PDF: pypdf
- Excel: openpyxl (.xlsx, .xlsm)
- Word: python-docx (.docx)
- HWP: graceful fallback (Hancom format — no pure Python parser; needs LibreOffice)
- Images: graceful fallback (would need OCR)
"""

import io
from pathlib import Path

# Hard limit on extracted text — keeps prompts cheap and prevents runaway costs.
MAX_CHARS = 30_000


def extract_text(filename: str, data: bytes) -> dict:
    """
    Returns: {ok: bool, text: str, kind: str, note: str}
    `kind` describes the source format. `note` is a short human-readable status.
    """
    name = (filename or "").lower()
    suffix = Path(name).suffix

    # --- Plain text-ish files ---
    text_suffixes = {".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml",
                     ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css",
                     ".sql", ".sh", ".env", ".ini", ".toml", ".rst", ".xml"}
    if suffix in text_suffixes:
        try:
            text = data.decode("utf-8", errors="replace")
            return _wrap(text, "text", f"Read {len(text)} chars as text")
        except Exception as e:
            return {"ok": False, "text": "", "kind": "text", "note": f"Decode error: {e}"}

    # --- PDF ---
    if suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(pages)
            return _wrap(text, "pdf", f"Extracted {len(reader.pages)} pages")
        except Exception as e:
            return {"ok": False, "text": "", "kind": "pdf", "note": f"PDF parse failed: {e}"}

    # --- Excel ---
    if suffix in (".xlsx", ".xlsm"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"=== Sheet: {ws.title} ===")
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c) for c in row]
                    if any(c for c in cells):
                        parts.append(" | ".join(cells))
                parts.append("")
            text = "\n".join(parts)
            return _wrap(text, "excel", f"Extracted {len(wb.worksheets)} sheets")
        except Exception as e:
            return {"ok": False, "text": "", "kind": "excel", "note": f"Excel parse failed: {e}"}

    # --- Word DOCX ---
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
            return _wrap(text, "docx", f"Extracted {len(paragraphs)} paragraphs")
        except Exception as e:
            return {"ok": False, "text": "", "kind": "docx", "note": f"DOCX parse failed: {e}"}

    # --- HWP (Korean Hancom) ---
    if suffix == ".hwp":
        return {
            "ok": False, "text": "", "kind": "hwp",
            "note": "HWP is not directly supported. Convert to PDF or DOCX in Hancom and re-upload.",
        }

    # --- Images ---
    if suffix in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
        return {
            "ok": False, "text": "", "kind": "image",
            "note": "Image OCR not yet implemented. Vision LLMs will be used in a future update.",
        }

    # --- Unknown — try as text ---
    try:
        text = data.decode("utf-8", errors="strict")
        return _wrap(text, "unknown-text", f"Decoded {len(text)} chars as fallback text")
    except UnicodeDecodeError:
        return {
            "ok": False, "text": "", "kind": "binary",
            "note": f"Binary file '{suffix}' — extraction not supported.",
        }


def _wrap(text: str, kind: str, note: str) -> dict:
    truncated = False
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + f"\n\n... [truncated at {MAX_CHARS} chars]"
        truncated = True
    return {
        "ok": bool(text.strip()),
        "text": text,
        "kind": kind,
        "note": note + (" (truncated)" if truncated else ""),
    }
